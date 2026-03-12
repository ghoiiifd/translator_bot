import os
import time
import telebot
import pdfplumber
from google import genai
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

TELEGRAM_BOT_TOKEN = os.environ.get('TELEGRAM_BOT_TOKEN', '8761170089:AAF3dOFS9BftYErhtUYiwdcbujkUBzhzgEs')
GEMINI_KEYS = [k for k in [os.environ.get('GEMINI_API_KEY', ''), os.environ.get('GEMINI_API_KEY_1', '')] if k]
if not GEMINI_KEYS:
            GEMINI_KEYS = ['AIzaSyAlxMiRdNW232qurvyZCvde_FbFilpUTao']

key_index = 0
bot = telebot.TeleBot(TELEGRAM_BOT_TOKEN)

def translate_text(text):
            global key_index
            if not text or len(text.strip()) < 3:
                            return ''
                        prompt = 'You are an expert academic translator. Translate from English to Arabic. Return only Arabic translation:\n\n' + text
    for _ in range(len(GEMINI_KEYS) * 2):
                    try:
                                        client = genai.Client(api_key=GEMINI_KEYS[key_index])
                                        response = client.models.generate_content(model='gemini-2.0-flash', contents=prompt)
                                        return response.text.strip()
except Exception as e:
            if '429' in str(e) or 'RESOURCE_EXHAUSTED' in str(e):
                                    key_index = (key_index + 1) % len(GEMINI_KEYS)
                                    time.sleep(2)
else:
                return '[\u062e\u0637\u0623 \u0641\u064a \u0627\u0644\u062a\u0631\u062c\u0645\u0629]'
            return '[\u062a\u062c\u0627\u0648\u0632 \u062d\u062f \u0627\u0644\u0637\u0644\u0628\u0627\u062a]'

def process_docx(inp, out):
            doc = Document(inp)
    new_doc = Document()
    for para in doc.paragraphs:
                    if para.text.strip():
                                        new_doc.add_paragraph(para.text)
                                        ar_para = new_doc.add_paragraph(translate_text(para.text))
                                        ar_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                        new_doc.add_paragraph('')
                                new_doc.save(out)

def process_pdf(inp, out):
            new_doc = Document()
                                with pdfplumber.open(inp) as pdf:
                                                for page in pdf.pages:
                                                                    text = page.extract_text()
                                                                    if text:
                                                                                            for block in text.split('\n\n'):
                                                                                                                        clean = block.replace('\n', ' ').strip()
                                                                                                                        if clean:
                                                                                                                                                        new_doc.add_paragraph(clean)
                                                                                                                                                        ar_para = new_doc.add_paragraph(translate_text(clean))
                                                                                                                                                        ar_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                                                                                                                                        new_doc.add_paragraph('')
                                                                                                                                    new_doc.save(out)
                                                                                                    
                                                                            @bot.message_handler(commands=['start', 'help'])
                                                        def welcome(message):
                                                                    bot.reply_to(message, '\u0645\u0631\u062d\u0628\u0627\u064b! \u0623\u0631\u0633\u0644 \u0645\u0644\u0641 PDF \u0623\u0648 DOCX \u0648\u0633\u0623\u062a\u0631\u062c\u0645\u0647 \u0644\u0643.')

                                        @bot.message_handler(content_types=['document'])
def handle_docs(message):
            ext = os.path.splitext(message.document.file_name)[1].lower()
    if ext not in ['.pdf', '.docx']:
                    bot.reply_to(message, '\u0623\u062f\u0639\u0645 \u0641\u0642\u0637 PDF \u0648 DOCX.')
        return
    bot.reply_to(message, '\u062c\u0627\u0631\u064a \u0627\u0644\u062a\u0631\u062c\u0645\u0629...')
    try:
                    data = bot.download_file(bot.get_file(message.document.file_id).file_path)
        inp = f'in_{message.chat.id}{ext}'
        out = f'out_{message.chat.id}.docx'
        open(inp, 'wb').write(data)
        if ext == '.docx':
                            process_docx(inp, out)
else:
            process_pdf(inp, out)
        bot.send_document(message.chat.id, open(out, 'rb'), caption='\u062a\u0645\u062a \u0627\u0644\u062a\u0631\u062c\u0645\u0629!')
        os.remove(inp)
        os.remove(out)
except Exception as e:
        bot.reply_to(message, f'\u062e\u0637\u0623: {e}')

if __name__ == '__main__':
            print('Bot running...')
    bot.infinity_polling()
